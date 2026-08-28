"""Compliance report generation: PDF (ReportLab), DOCX (python-docx), CSV."""
import os, json, csv
from datetime import datetime
from . import db

COLORS = {"PASS": "#1a7f37", "FAIL": "#cf222e", "WARN": "#9a6700", "INFO": "#0969da",
          "NON-COMPLIANT": "#cf222e", "PARTIALLY COMPLIANT": "#9a6700", "COMPLIANT": "#1a7f37",
          "CRITICAL": "#b60205", "MAJOR": "#d73a49", "MINOR": "#9a6700", "OK": "#1a7f37", "INFO2": "#0969da"}

PENALTY_REF = [
    ("Sec. 35, LMA 2009", "Contravention of s.18 (declarations on pre-packaged commodities)",
     "Fine up to ₹25,000 (1st); ₹50,000 (2nd); ₹1,00,000 (3rd & subsequent)"),
    ("Sec. 36(1), LMA 2009", "Manufacturing/packing/selling non-standard packages (non-conforming declarations)",
     "Fine up to ₹25,000 (1st); ₹50,000 (2nd); ≥₹50,000–₹1,00,000 or imprisonment up to 1 yr or both (subsequent)"),
    ("Sec. 36(2), LMA 2009", "Pre-packed commodity with error in net quantity",
     "Fine ₹10,000–₹50,000 (1st); up to ₹1,00,000 or imprisonment up to 1 yr or both (2nd & subsequent)"),
    ("Sec. 39, LMA 2009", "General penalty for contravention of rules",
     "Fine up to ₹25,000 (1st); ₹50,000 (2nd); ₹1,00,000 (3rd & subsequent)"),
]

def _load(scan_id):
    conn = db.get_conn()
    row = conn.execute("SELECT * FROM scans WHERE id=?", (scan_id,)).fetchone()
    conn.close()
    return row

def _meta(row):
    checks = json.loads(row["checks_json"] or "[]")
    fields = json.loads(row["font_json"] or "{}")
    ext = json.loads(row["extracted_json"] or "{}")
    return checks, fields, ext

# ------------------------------------------------------------------ PDF
def pdf_report(scan_id, out_dir, created_by=None):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable)
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle

    row = _load(scan_id)
    if not row:
        return None
    checks, fields, ext = _meta(row)
    path = os.path.join(out_dir, f"LMR_{row['scan_code']}_report.pdf")
    st = getSampleStyleSheet()
    title = ParagraphStyle("t", parent=st["Title"], fontSize=16, textColor=colors.HexColor("#0b3d66"))
    sub = ParagraphStyle("s", parent=st["Normal"], fontSize=9, textColor=colors.HexColor("#555"))
    h2 = ParagraphStyle("h2", parent=st["Heading2"], fontSize=11, spaceBefore=10, spaceAfter=4,
                        textColor=colors.HexColor("#0b3d66"))
    cell = ParagraphStyle("c", parent=st["Normal"], fontSize=8, leading=11)
    cellb = ParagraphStyle("cb", parent=cell, fontName="Helvetica-Bold")

    doc = SimpleDocTemplate(path, pagesize=A4, topMargin=18 * mm, bottomMargin=16 * mm,
                            leftMargin=16 * mm, rightMargin=16 * mm,
                            title=f"Legal Metrology Compliance Report {row['scan_code']}")
    el = []
    el.append(Paragraph("LEGAL METROLOGY COMPLIANCE REPORT", title))
    el.append(Paragraph("Automated label / listing compliance check under the Legal Metrology "
                        "(Packaged Commodities) Rules, 2011 (as amended)", sub))
    el.append(Spacer(1, 4))
    el.append(HRFlowable(width="100%", thickness=1.2, color=colors.HexColor("#0b3d66")))
    el.append(Spacer(1, 8))

    status_col = COLORS.get(row["overall_status"], "#333")
    meta = [
        [Paragraph("Scan Code", cellb), Paragraph(str(row["scan_code"]), cell),
         Paragraph("Overall Status", cellb),
         Paragraph(f"<b><font color='{status_col}'>{row['overall_status']}</font></b> "
                   f"&nbsp;&nbsp;Score: <b>{row['score']}/100</b>", cell)],
        [Paragraph("Product", cellb), Paragraph(str(row["product_name"]), cell),
         Paragraph("Category", cellb), Paragraph(str(row["category"] or "—"), cell)],
        [Paragraph("Source", cellb), Paragraph("Package image" if row["source"] == "image" else "E-commerce listing", cell),
         Paragraph("Scanned on", cellb), Paragraph(str(row["created_at"]), cell)],
        [Paragraph("PDP area", cellb), Paragraph(f"{row['pdp_area_cm2']:g} cm²" if row["pdp_area_cm2"] else "n/a", cell),
         Paragraph("Review status", cellb), Paragraph(str(row["review_status"]), cell)],
    ]
    t = Table(meta, colWidths=[32 * mm, 62 * mm, 32 * mm, 66 * mm])
    t.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#bbb")),
                           ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#eef3f8")),
                           ("BACKGROUND", (2, 0), (2, -1), colors.HexColor("#eef3f8")),
                           ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    el.append(t)
    el.append(Spacer(1, 8))

    # ---- detected declarations
    el.append(Paragraph("1. Declarations Detected / Extracted", h2))
    det_rows = [[Paragraph("Mandatory declaration", cellb), Paragraph("Detected on label/listing", cellb)]]
    dd = ext
    det_rows.append([Paragraph("Name & address of manufacturer/packer/importer (R6(1)(a))", cell),
                     Paragraph(dd.get("address", {}).get("snippet") or "NOT DETECTED", cell)])
    det_rows.append([Paragraph("Country of origin (R6(1)(aa))", cell),
                     Paragraph((dd.get("country", {}).get("origin") or "—") +
                               (" (imported)" if dd.get("country", {}).get("imported") else ""), cell)])
    det_rows.append([Paragraph("Common/generic name (R6(1)(b))", cell),
                     Paragraph("detected" if dd.get("common_name", {}).get("found") else "NOT DETECTED", cell)])
    nq = (dd.get("net_qty") or [{}])[0]
    det_rows.append([Paragraph("Net quantity (R6(1)(c))", cell),
                     Paragraph(f"{nq.get('value','—')} {nq.get('unit','')} ({nq.get('type','')})" if nq else "NOT DETECTED", cell)])
    mf = (dd.get("mfg_date") or [{}])[0]
    det_rows.append([Paragraph("Month & year of manufacture (R6(1)(d))", cell),
                     Paragraph(mf.get("matched", "NOT DETECTED"), cell)])
    bb = (dd.get("best_before") or [{}])[0]
    det_rows.append([Paragraph("Best before / use by (R6(1)(da))", cell),
                     Paragraph(bb.get("matched", "—"), cell)])
    mr = (dd.get("mrp") or [{}])[0]
    det_rows.append([Paragraph("Retail sale price / MRP (R6(1)(e))", cell),
                     Paragraph(f"Rs. {mr['value']:.2f}" + (" (incl. of all taxes)" if mr.get("incl_taxes") else "") if mr else "NOT DETECTED", cell)])
    cc = dd.get("consumer_care", {})
    det_rows.append([Paragraph("Consumer care details (R6(2))", cell),
                     Paragraph(", ".join(cc.get("phones", []) + cc.get("emails", [])) or "NOT DETECTED", cell)])
    t2 = Table(det_rows, colWidths=[80 * mm, 112 * mm])
    t2.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#bbb")),
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0b3d66")),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                            ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    el.append(t2)
    el.append(Spacer(1, 8))

    # ---- checks
    el.append(Paragraph("2. Compliance Checks & Violations", h2))
    ch_rows = [[Paragraph("#", cellb), Paragraph("Check (Rule reference)", cellb),
                Paragraph("Status", cellb), Paragraph("Severity", cellb), Paragraph("Finding / Evidence", cellb)]]
    for i, c in enumerate(checks, 1):
        st_col = colors.HexColor(COLORS.get(c["status"], "#333"))
        sv_col = colors.HexColor(COLORS.get(c["severity"], "#333"))
        ch_rows.append([Paragraph(str(i), cell),
                        Paragraph(f"<b>{c['title']}</b><br/><font size=7 color='#666'>{c['rule']}</font>", cell),
                        Paragraph(f"<font color='{st_col.hexval()}'><b>{c['status']}</b></font>", cell),
                        Paragraph(f"<font color='{sv_col.hexval()}'>{c['severity']}</font>", cell),
                        Paragraph(c["message"] + (f"<br/><i>Evidence: {c['evidence']}</i>" if c.get("evidence") else ""), cell)])
    t3 = Table(ch_rows, colWidths=[8 * mm, 44 * mm, 20 * mm, 20 * mm, 100 * mm], repeatRows=1)
    t3.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#bbb")),
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0b3d66")),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                            ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    el.append(t3)
    el.append(Spacer(1, 8))

    # ---- font sizes
    if fields:
        el.append(Paragraph("3. Font Size & Readability Analysis (Rule 7 Table-I)", h2))
        req = None
        try:
            from .fontsize import required_height_mm
            req = required_height_mm(row["pdp_area_cm2"] or 0)
        except Exception:
            req = None
        fs_rows = [[Paragraph("Declaration field", cellb), Paragraph("Measured height (mm)", cellb),
                    Paragraph("Required (mm)", cellb), Paragraph("Status", cellb)]]
        for k, f in fields.items():
            ok = req is None or f["height_mm"] >= req - 0.15
            fs_rows.append([Paragraph(f["label"], cell),
                            Paragraph(f"{f['height_mm']:.2f}", cell),
                            Paragraph(f"{req:.1f}" if req else "—", cell),
                            Paragraph(f"<font color='{'#1a7f37' if ok else '#cf222e'}'><b>{'OK' if ok else 'TOO SMALL'}</b></font>", cell)])
        t4 = Table(fs_rows, colWidths=[100 * mm, 40 * mm, 32 * mm, 40 * mm], repeatRows=1)
        t4.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#bbb")),
                                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0b3d66")),
                                ("TEXTCOLOR", (0, 0), (-1, 0), colors.white)]))
        el.append(t4)
        el.append(Spacer(1, 8))

    # ---- penalty reference
    el.append(Paragraph("4. Penalty Reference (Legal Metrology Act, 2009)", h2))
    p_rows = [[Paragraph("Provision", cellb), Paragraph("Offence", cellb), Paragraph("Penalty", cellb)]]
    for a, b, c in PENALTY_REF:
        p_rows.append([Paragraph(a, cell), Paragraph(b, cell), Paragraph(c, cell)])
    t5 = Table(p_rows, colWidths=[40 * mm, 78 * mm, 74 * mm], repeatRows=1)
    t5.setStyle(TableStyle([("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#bbb")),
                            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#0b3d66")),
                            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
                            ("VALIGN", (0, 0), (-1, -1), "TOP")]))
    el.append(t5)
    el.append(Spacer(1, 10))

    el.append(Paragraph("Notes: " + (row["notes"] or "—"), cell))
    el.append(Spacer(1, 16))
    el.append(HRFlowable(width="100%", thickness=0.7, color=colors.HexColor("#999")))
    el.append(Spacer(1, 10))
    sig = Table([[Paragraph("Inspector / Enforcement Officer", cellb),
                  Paragraph("Signature", cellb), Paragraph("Date", cellb)],
                 [Paragraph("Name: ____________", cell), Paragraph("_____________", cell),
                  Paragraph(datetime.now().strftime("%d-%m-%Y"), cell)]],
                colWidths=[64 * mm, 64 * mm, 64 * mm])
    sig.setStyle(TableStyle([("BOX", (0, 0), (-1, -1), 0.5, colors.HexColor("#999")),
                             ("INNERGRID", (0, 0), (-1, -1), 0.3, colors.HexColor("#ccc")),
                             ("TOPPADDING", (0, 0), (-1, -1), 10)]))
    el.append(sig)
    el.append(Spacer(1, 8))
    el.append(Paragraph("This report is an automated screening aid generated by LM Compliance System. "
                        "Findings should be verified physically by a Legal Metrology Inspector before any "
                        "enforcement action. Legal references: Legal Metrology Act, 2009; Legal Metrology "
                        "(Packaged Commodities) Rules, 2011 with amendments up to GSR 629(E) dt. 23.06.2017.", sub))
    doc.build(el)
    return path

# ------------------------------------------------------------------ DOCX
def docx_report(scan_id, out_dir):
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    row = _load(scan_id)
    if not row:
        return None
    checks, fields, ext = _meta(row)
    path = os.path.join(out_dir, f"LMR_{row['scan_code']}_report.docx")
    doc = Document()
    doc.add_heading("Legal Metrology Compliance Report", level=0)
    p = doc.add_paragraph("Automated label / listing compliance check under the Legal Metrology "
                          "(Packaged Commodities) Rules, 2011 (as amended)")
    p.runs[0].italic = True
    meta = doc.add_table(rows=4, cols=4)
    meta.style = "Light Grid Accent 1"
    vals = [
        ("Scan Code", row["scan_code"], "Overall Status", row["overall_status"] + f"  (Score: {row['score']}/100)"),
        ("Product", row["product_name"], "Category", row["category"] or "—"),
        ("Source", "Package image" if row["source"] == "image" else "E-commerce listing",
         "Scanned on", row["created_at"]),
        ("PDP area", f"{row['pdp_area_cm2']:g} cm²" if row["pdp_area_cm2"] else "n/a",
         "Review status", row["review_status"]),
    ]
    for i, (a, b, c, d) in enumerate(vals):
        for j, v in enumerate((a, b, c, d)):
            cell = meta.cell(i, j)
            cell.text = str(v)
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)

    doc.add_heading("1. Declarations Detected", level=1)
    t = doc.add_table(rows=9, cols=2)
    t.style = "Light Grid Accent 1"
    dd = ext
    rows = [
        ("Name & address of manufacturer/packer/importer (R6(1)(a))", dd.get("address", {}).get("snippet") or "NOT DETECTED"),
        ("Country of origin (R6(1)(aa))", (dd.get("country", {}).get("origin") or "—") +
         (" (imported)" if dd.get("country", {}).get("imported") else "")),
        ("Common/generic name (R6(1)(b))", "detected" if dd.get("common_name", {}).get("found") else "NOT DETECTED"),
        ("Net quantity (R6(1)(c))", (lambda nq: f"{nq.get('value','—')} {nq.get('unit','')}" if nq else "NOT DETECTED")((dd.get("net_qty") or [{}])[0])),
        ("Month & year of manufacture (R6(1)(d))", (dd.get("mfg_date") or [{}])[0].get("matched", "NOT DETECTED")),
        ("Best before / use by (R6(1)(da))", (dd.get("best_before") or [{}])[0].get("matched", "—")),
        ("Retail sale price / MRP (R6(1)(e))", (lambda m: f"Rs. {m['value']:.2f}" if m else "NOT DETECTED")((dd.get("mrp") or [{}])[0])),
        ("Consumer care details (R6(2))", ", ".join(dd.get("consumer_care", {}).get("phones", []) +
                                                    dd.get("consumer_care", {}).get("emails", [])) or "NOT DETECTED"),
    ]
    for i, (a, b) in enumerate(rows):
        t.cell(i, 0).text, t.cell(i, 1).text = a, str(b)
        for cell in (t.cell(i, 0), t.cell(i, 1)):
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(9)

    doc.add_heading("2. Compliance Checks", level=1)
    t2 = doc.add_table(rows=len(checks) + 1, cols=4)
    t2.style = "Light Grid Accent 1"
    for j, h in enumerate(["Check (Rule)", "Status", "Severity", "Finding"]):
        t2.cell(0, j).text = h
    for i, c in enumerate(checks, 1):
        t2.cell(i, 0).text = f"{c['title']} [{c['rule']}]"
        t2.cell(i, 1).text = c["status"]
        t2.cell(i, 2).text = c["severity"]
        t2.cell(i, 3).text = c["message"] + (f" | Evidence: {c['evidence']}" if c.get("evidence") else "")
        for j in range(4):
            for run in t2.cell(i, j).paragraphs[0].runs:
                run.font.size = Pt(8)

    if fields:
        doc.add_heading("3. Font Size Analysis (Rule 7 Table-I)", level=1)
        t3 = doc.add_table(rows=len(fields) + 1, cols=4)
        t3.style = "Light Grid Accent 1"
        for j, h in enumerate(["Field", "Measured (mm)", "Required (mm)", "Status"]):
            t3.cell(0, j).text = h
        try:
            from .fontsize import required_height_mm
            req = required_height_mm(row["pdp_area_cm2"] or 0)
        except Exception:
            req = None
        for i, (k, f) in enumerate(fields.items(), 1):
            ok = req is None or f["height_mm"] >= req - 0.15
            t3.cell(i, 0).text = f["label"]
            t3.cell(i, 1).text = f"{f['height_mm']:.2f}"
            t3.cell(i, 2).text = f"{req:.1f}" if req else "—"
            t3.cell(i, 3).text = "OK" if ok else "TOO SMALL"

    doc.add_heading("4. Penalty Reference", level=1)
    t4 = doc.add_table(rows=len(PENALTY_REF) + 1, cols=3)
    t4.style = "Light Grid Accent 1"
    for j, h in enumerate(["Provision", "Offence", "Penalty"]):
        t4.cell(0, j).text = h
    for i, (a, b, c) in enumerate(PENALTY_REF, 1):
        t4.cell(i, 0).text, t4.cell(i, 1).text, t4.cell(i, 2).text = a, b, c

    doc.add_paragraph()
    doc.add_paragraph(f"Notes: {row['notes'] or '—'}")
    doc.add_paragraph(f"Inspector: ____________    Signature: ____________    Date: {datetime.now().strftime('%d-%m-%Y')}")
    doc.add_paragraph("This report is an automated screening aid; verify findings physically before "
                      "enforcement action.").runs[0].italic = True
    doc.save(path)
    return path

# ------------------------------------------------------------------ CSV
def csv_report(scan_id, out_dir):
    row = _load(scan_id)
    if not row:
        return None
    checks, fields, _ = _meta(row)
    path = os.path.join(out_dir, f"LMR_{row['scan_code']}_checks.csv")
    with open(path, "w", newline="", encoding="utf-8") as f:
        w = csv.writer(f)
        w.writerow(["scan_code", "product", "category", "source", "overall_status", "score",
                    "check_id", "rule_ref", "check", "status", "severity", "finding", "evidence", "suggestion"])
        for c in checks:
            w.writerow([row["scan_code"], row["product_name"], row["category"], row["source"],
                        row["overall_status"], row["score"], c["id"], c["rule"], c["title"],
                        c["status"], c["severity"], c["message"], c.get("evidence", ""), c.get("suggestion", "")])
    return path
